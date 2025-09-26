import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import re
import io
from datetime import datetime
import random

# ====== CONFIGURACIÓN DE LA PÁGINA ======
st.set_page_config(
    page_title="🛡️ Generador de Certificados Pacífico Seguros",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ====== CSS PERSONALIZADO ======
st.markdown("""
<style>
    /* Ocultar elementos de Streamlit */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Estilo del encabezado */
    .main-header {
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        padding: 2rem;
        border-radius: 10px;
        text-align: center;
        margin-bottom: 2rem;
    }
    
    /* Estilo de las tarjetas de progreso */
    .progress-card {
        background: #f8f9ff;
        border-left: 4px solid #1e3c72;
        padding: 1rem;
        margin: 1rem 0;
        border-radius: 5px;
    }
    
    /* Estilo de los campos de entrada */
    .stTextInput > div > div > input {
        border-radius: 5px;
        border: 2px solid #e0e0e0;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #1e3c72;
        box-shadow: 0 0 0 0.2rem rgba(30, 60, 114, 0.25);
    }
</style>
""", unsafe_allow_html=True)

# ====== CLASE PARA GENERAR EL CERTIFICADO ======
class CertificadoPacificoGenerator:
    def __init__(self):
        self.variables = [
            'nombre_seguro',
            'numero_certificado', 
            'numero_poliza',
            'codigo_registro',
            'nombre_asegurado',
            'riesgo_protegido',
            'fecha_inicio',
            'tipo_documento',
            'numero_documento',
            'domicilio_asegurado',
            'correo_asegurado',
            'telefono_asegurado',
            'relacion_contratante',
            'situaciones_cobertura',
            'situaciones_adicionales',
            'condiciones_asegurabilidad',
            'situaciones_no_cubiertas',
            'uso_cobertura',
            'moneda',
            'costo_total',
            'igv',
            'frecuencia_pago',
            'medio_cobro',
            'duracion_seguro',
            'requisitos_vigencia',
            'condiciones_fin',
            'fecha_emision'
        ]
        
        self.questions = {
            'nombre_seguro': {
                'question': '¿Cuál es el nombre del seguro?',
                'help': 'Ejemplo: Seguro de Vida, Seguro Vehicular, etc.',
                'type': 'text'
            },
            'numero_certificado': {
                'question': '¿Cuál es el número de certificado?',
                'help': 'Número único del certificado',
                'type': 'text'
            },
            'numero_poliza': {
                'question': '¿Cuál es el número de póliza?',
                'help': 'Número de la póliza de seguro',
                'type': 'text'
            },
            'codigo_registro': {
                'question': '¿Cuál es el código de registro?',
                'help': 'Código interno de registro',
                'type': 'text'
            },
            'nombre_asegurado': {
                'question': '¿Cuál es el nombre completo del asegurado?',
                'help': 'Nombres y apellidos completos',
                'type': 'text'
            },
            'riesgo_protegido': {
                'question': '¿Contra qué riesgo está protegido?',
                'help': 'Ejemplo: fallecimiento, accidentes, robo, etc.',
                'type': 'text'
            },
            'fecha_inicio': {
                'question': '¿Cuál es la fecha de inicio de vigencia?',
                'help': 'Fecha en que inicia la cobertura',
                'type': 'date'
            },
            'tipo_documento': {
                'question': '¿Qué tipo de documento de identidad?',
                'help': 'DNI, Pasaporte, Carnet de Extranjería, etc.',
                'type': 'text'
            },
            'numero_documento': {
                'question': '¿Cuál es el número de documento?',
                'help': 'Número del documento de identidad',
                'type': 'text'
            },
            'domicilio_asegurado': {
                'question': '¿Cuál es el domicilio del asegurado?',
                'help': 'Dirección completa',
                'type': 'text'
            },
            'correo_asegurado': {
                'question': '¿Cuál es el correo electrónico del asegurado?',
                'help': 'Email de contacto',
                'type': 'email'
            },
            'telefono_asegurado': {
                'question': '¿Cuál es el teléfono del asegurado?',
                'help': 'Número de teléfono',
                'type': 'text'
            },
            'relacion_contratante': {
                'question': '¿Cuál es la relación del asegurado con el contratante?',
                'help': 'Ejemplo: Titular, Beneficiario, Cónyuge, etc.',
                'type': 'text'
            },
            'situaciones_cobertura': {
                'question': '¿En qué situaciones cubre el seguro?',
                'help': 'Describe las coberturas principales',
                'type': 'textarea'
            },
            'situaciones_adicionales': {
                'question': '¿Qué situaciones adicionales cubre?',
                'help': 'Coberturas adicionales o complementarias',
                'type': 'textarea'
            },
            'condiciones_asegurabilidad': {
                'question': '¿Cuáles son las condiciones de asegurabilidad?',
                'help': 'Condiciones especiales que debe cumplir',
                'type': 'textarea'
            },
            'situaciones_no_cubiertas': {
                'question': '¿Qué situaciones NO cubre el seguro?',
                'help': 'Exclusiones de la póliza',
                'type': 'textarea'
            },
            'uso_cobertura': {
                'question': '¿Cómo hacer uso de la cobertura?',
                'help': 'Procedimiento para hacer un reclamo',
                'type': 'textarea'
            },
            'moneda': {
                'question': '¿En qué moneda está el seguro?',
                'help': 'Soles, Dólares, etc.',
                'type': 'text'
            },
            'costo_total': {
                'question': '¿Cuál es el costo total del seguro?',
                'help': 'Monto total sin IGV',
                'type': 'number'
            },
            'igv': {
                'question': '¿Cuál es el monto del IGV?',
                'help': 'Impuesto General a las Ventas',
                'type': 'number'
            },
            'frecuencia_pago': {
                'question': '¿Cuál es la frecuencia de pago?',
                'help': 'Mensual, Anual, Trimestral, etc.',
                'type': 'text'
            },
            'medio_cobro': {
                'question': '¿Cómo se cobra el seguro?',
                'help': 'Tarjeta de crédito, débito automático, etc.',
                'type': 'text'
            },
            'duracion_seguro': {
                'question': '¿Cuánto dura el seguro?',
                'help': 'Un mes, un año, etc.',
                'type': 'text'
            },
            'requisitos_vigencia': {
                'question': '¿Qué requisitos debe cumplir para que empiece la vigencia?',
                'help': 'Condiciones para que inicie la cobertura',
                'type': 'textarea'
            },
            'condiciones_fin': {
                'question': '¿En qué condiciones termina el seguro?',
                'help': 'Situaciones que dan fin al contrato',
                'type': 'textarea'
            },
            'fecha_emision': {
                'question': '¿Cuál es la fecha de emisión del certificado?',
                'help': 'Fecha en que se emite este certificado',
                'type': 'date'
            }
        }
    
    def generate_certificate_document(self, responses):
        """Genera el documento Word del certificado con el formato original"""
        
        # Crear nuevo documento
        doc = Document()
        
        # Configurar márgenes y formato general
        section = doc.sections[0]
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # ENCABEZADO
        header = doc.add_paragraph()
        header_run = header.add_run(f"Certificado N° {responses.get('numero_certificado', 'XXXXXXX')} -- Seguro de ")
        header_run.italic = True
        
        header_run2 = header.add_run(responses.get('nombre_seguro', '[NOMBRE DEL SEGURO]'))
        header_run2.bold = True
        
        # Número de póliza
        poliza_p = doc.add_paragraph()
        poliza_run = poliza_p.add_run(f"Póliza Nº {responses.get('numero_poliza', 'XXXXX')} - Código de registro {responses.get('codigo_registro', 'XXXXXXX')}")
        poliza_run.italic = True
        
        # Saludo
        doc.add_paragraph(f"¡Hola {responses.get('nombre_asegurado', 'XXXXXXXXX')}!")
        doc.add_paragraph("¡Felicidades! Estás asegurado.")
        
        # Confirmación de seguro
        confirmacion = doc.add_paragraph("Confirmamos que tienes un seguro activo que te protege frente a ")
        confirmacion_run = confirmacion.add_run(responses.get('riesgo_protegido', '[COMPLETAR CON EL RIESGO]'))
        confirmacion_run.bold = True
        
        # CONTRATANTE
        contratante_title = doc.add_paragraph()
        contratante_run = contratante_title.add_run("CONTRATANTE")
        contratante_run.bold = True
        
        doc.add_paragraph("XXXXX, RUC xxxxxxx, Dirección xxxxxxxxx")
        doc.add_paragraph("Distrito xxxxxxx xxxxxxx también llamado sólo \"xxxxx\".")
        
        # VIGENCIA
        vigencia_title = doc.add_paragraph()
        vigencia_run = vigencia_title.add_run("Vigencia del Seguro: XXXXXXXXXXX")
        vigencia_run.bold = True
        
        doc.add_paragraph(f"Inicio de Vigencia: Desde las XX horas del {responses.get('fecha_inicio', 'DD/MM/AAA')}")
        doc.add_paragraph("Tu seguro se renovará automáticamente.")
        
        # INFORMACIÓN DE CONTACTO
        contacto_title = doc.add_paragraph()
        contacto_run = contacto_title.add_run("Información de Contacto de Pacífico Seguros")
        contacto_run.bold = True
        
        doc.add_paragraph("Pacífico Compañía de Seguros y Reaseguros S.A.")
        doc.add_paragraph("RUC N 20332970411 Av. Juan de Arona 830, San Isidro")
        doc.add_paragraph("Teléf.: XXX-XXXX / WhatsApp: +51 XXX-XXXX")
        doc.add_paragraph("Pág. Web.: https://www.pacifico.com.pe/")
        
        # Mensaje de contacto
        contacto_msg = doc.add_paragraph()
        contacto_msg_run = contacto_msg.add_run("Si tienes alguna duda sobre tu cobertura o cómo usar tu seguro, contáctanos al número de teléfono indicado o escríbenos por WhatsApp.")
        contacto_msg_run.bold = True
        
        # DATOS DEL ASEGURADO
        asegurado_title = doc.add_paragraph()
        asegurado_run = asegurado_title.add_run("¿Quién es el ASEGURADO?")
        asegurado_run.bold = True
        
        doc.add_paragraph(responses.get('nombre_asegurado', '[Nombre y Apellidos del Asegurado]'))
        doc.add_paragraph("¡Tú estás asegurado!")
        doc.add_paragraph(responses.get('tipo_documento', '[Tipo Doc]'))
        doc.add_paragraph(responses.get('numero_documento', '[Número Doc]'))
        doc.add_paragraph(responses.get('domicilio_asegurado', '[Domicilio]'))
        doc.add_paragraph(responses.get('correo_asegurado', '[Correo]'))
        doc.add_paragraph(responses.get('telefono_asegurado', '[Teléfono]'))
        
        # Domicilio contractual
        domicilio_contractual = doc.add_paragraph()
        domicilio_contractual_run = domicilio_contractual.add_run("Tu domicilio contractual será el correo electrónico que brindaste en la Solicitud de Seguro. Si no lo hiciste, será la dirección física ingresada en los sistemas del [completar con la info del canal. Por ejemplo, para PT es el \"Banco\"].")
        domicilio_contractual_run.bold = True
        
        doc.add_paragraph(f"Relación del ASEGURADO con el CONTRATANTE: {responses.get('relacion_contratante', 'XXXXXXX')}")
        
        # COBERTURAS
        cobertura_title = doc.add_paragraph()
        cobertura_run = cobertura_title.add_run("¿En qué situaciones te cubre tu seguro?")
        cobertura_run.bold = True
        
        cobertura_content = doc.add_paragraph()
        cobertura_content_run = cobertura_content.add_run(responses.get('situaciones_cobertura', '[Aquí debes modificar en función a los inputs]'))
        cobertura_content_run.bold = True
        
        # COBERTURAS ADICIONALES
        adicional_title = doc.add_paragraph()
        adicional_run = adicional_title.add_run("¿En qué situaciones adicionales te cubre tu seguro?")
        adicional_run.bold = True
        
        adicional_content = doc.add_paragraph()
        adicional_content_run = adicional_content.add_run(responses.get('situaciones_adicionales', 'xxxxxxxxxxxxxx'))
        adicional_content_run.bold = True
        
        # INFORMACIÓN IMPORTANTE
        info_title = doc.add_paragraph()
        info_run = info_title.add_run("¿Qué información importante debes considerar?")
        info_run.bold = True
        
        info_content = doc.add_paragraph()
        info_content_run = info_content.add_run(responses.get('condiciones_asegurabilidad', '[Completar con las condiciones de asegurabilidad]'))
        info_content_run.bold = True
        
        # EXCLUSIONES
        exclusiones_title = doc.add_paragraph()
        exclusiones_run = exclusiones_title.add_run("¿En qué situaciones que NO cubre tu seguro?")
        exclusiones_run.bold = True
        
        exclusiones_content = doc.add_paragraph()
        exclusiones_content_run = exclusiones_content.add_run(responses.get('situaciones_no_cubiertas', '[Aquí debes modificar en función a los inputs]'))
        exclusiones_content_run.bold = True
        
        # USO DE COBERTURA
        uso_title = doc.add_paragraph()
        uso_run = uso_title.add_run("¿Cómo hago uso de la cobertura?")
        uso_run.bold = True
        
        uso_intro = doc.add_paragraph()
        uso_intro_run = uso_intro.add_run("Si sucediera alguna de las situaciones cubiertas por el seguro que describimos anteriormente:")
        uso_intro_run.bold = True
        
        uso_content = doc.add_paragraph()
        uso_content_run = uso_content.add_run(responses.get('uso_cobertura', '[Aquí debes modificar en función a los inputs]'))
        uso_content_run.bold = True
        
        # Límite de tiempo
        limite = doc.add_paragraph()
        limite_run = limite.add_run("El límite de tiempo que tienes para presentar tus documentos es de 10 años.")
        limite_run.bold = True
        
        # COSTOS
        costo_title = doc.add_paragraph()
        costo_run = costo_title.add_run("¿Cuánto cuesta y cómo se paga el seguro?")
        costo_run.bold = True
        
        # Tabla de costos
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        # Llenar tabla
        cells = table.rows[0].cells
        cells[0].text = "Moneda"
        cells[1].text = responses.get('moneda', 'xxxxxxx')
        
        cells = table.rows[1].cells
        cells[0].text = "Costo Total del Seguro"
        cells[1].text = str(responses.get('costo_total', 'xxxx'))
        
        cells = table.rows[2].cells
        cells[0].text = "IGV"
        cells[1].text = str(responses.get('igv', 'xxxx'))
        
        cells = table.rows[3].cells
        cells[0].text = "Frecuencia"
        cells[1].text = responses.get('frecuencia_pago', 'xxxx')
        
        cells = table.rows[4].cells
        cells[0].text = "¿Cómo te cobramos el seguro?"
        cells[1].text = responses.get('medio_cobro', '[completar la información del medio de cobro]')
        
        # DURACIÓN
        duracion_title = doc.add_paragraph()
        duracion_run = duracion_title.add_run("¿Cuánto dura tu seguro?")
        duracion_run.bold = True
        
        doc.add_paragraph(f"Tu seguro puede durar {responses.get('duracion_seguro', 'un mes o un año')}, según el plan que elegiste.")
        
        # INICIO Y FIN
        inicio_title = doc.add_paragraph()
        inicio_run = inicio_title.add_run("¿Cuándo empieza y cuándo termina?")
        inicio_run.bold = True
        
        inicio_subtitle = doc.add_paragraph()
        inicio_subtitle_run = inicio_subtitle.add_run("Inicio: Tu seguro empieza desde que lo contratas, si:")
        inicio_subtitle_run.bold = True
        
        inicio_content = doc.add_paragraph()
        inicio_content_run = inicio_content.add_run(responses.get('requisitos_vigencia', '[Completar con los requisitos propios del producto]'))
        inicio_content_run.bold = True
        
        fin_subtitle = doc.add_paragraph()
        fin_subtitle_run = fin_subtitle.add_run("Fin: Tu seguro terminará si ocurre alguna de estas situaciones:")
        fin_subtitle_run.bold = True
        
        fin_content = doc.add_paragraph()
        fin_content_run = fin_content.add_run(responses.get('condiciones_fin', '[Completar con las condiciones de fin]'))
        fin_content_run.bold = True
        
        # DERECHO DE ARREPENTIMIENTO
        arrepentimiento_title = doc.add_paragraph()
        arrepentimiento_run = arrepentimiento_title.add_run("¿Puedo arrepentirme de haber contratado el seguro?")
        arrepentimiento_run.bold = True
        
        doc.add_paragraph("Sí. Si cambias de opinión, puedes cancelar el seguro sin dar una razón y sin penalidades dentro de los 15 días calendario desde que recibiste este Certificado.")
        
        # FECHA DE EMISIÓN
        doc.add_paragraph("")
        fecha_emision = doc.add_paragraph()
        fecha_emision_run = fecha_emision.add_run(f"Fecha de emisión, Lima, {responses.get('fecha_emision', 'xx de xxxx de xxxx')}")
        fecha_emision_run.bold = True
        
        # FIRMA
        doc.add_paragraph("")
        firma = doc.add_paragraph()
        firma_run = firma.add_run("Xxxxxxxxxxxxxxxxxxxxxxxxxxx")
        firma_run.bold = True
        
        representante = doc.add_paragraph()
        representante_run = representante.add_run("Representante Pacífico Seguros")
        representante_run.bold = True
        
        return doc

# ====== FUNCIÓN PRINCIPAL ======
def main():
    # Inicializar estado
    if "step" not in st.session_state:
        st.session_state.step = "welcome"
        st.session_state.generator = CertificadoPacificoGenerator()
        st.session_state.current_question = 0
        st.session_state.responses = {}
        st.session_state.chat_history = []
    
    generator = st.session_state.generator
    
    # HEADER PRINCIPAL
    st.markdown("""
    <div class="main-header">
        <h1>🛡️ Generador de Certificados Pacífico Seguros</h1>
        <p>Asistente inteligente para generar certificados de seguro personalizados</p>
    </div>
    """, unsafe_allow_html=True)
    
    # NAVEGACIÓN
    if st.session_state.step == "welcome":
        show_welcome()
    elif st.session_state.step == "interview":
        show_interview()
    elif st.session_state.step == "generate":
        show_generation()

def show_welcome():
    """Pantalla de bienvenida"""
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("""
        ### 👋 Bienvenido al Generador de Certificados Pacífico Seguros
        
        Este asistente te ayudará a generar certificados de seguro personalizados con el formato oficial de Pacífico Seguros.
        
        **¿Cómo funciona?**
        
        1. 📝 **Responde las preguntas** sobre el seguro y el asegurado
        2. 🤖 **El asistente recopila** toda la información necesaria
        3. 📄 **Genera el certificado** con el formato oficial de Pacífico
        4. ⬇️ **Descarga tu documento** listo para usar
        
        **Características:**
        - ✅ **Formato oficial** de Pacífico Seguros
        - ✅ **Todos los campos** completados automáticamente  
        - ✅ **Conserva el diseño** original del certificado
        - ✅ **Listo para imprimir** o enviar por email
        """)
        
        if st.button("🚀 Comenzar a generar certificado", use_container_width=True):
            st.session_state.step = "interview"
            st.session_state.chat_history = [{
                "role": "assistant",
                "content": "¡Hola! Soy tu asistente para generar certificados de Pacífico Seguros. Te haré algunas preguntas para completar toda la información necesaria. ¿Empezamos? 🛡️"
            }]
            st.rerun()
    
    with col2:
        st.markdown("""
        <div class="progress-card">
            <h4>📊 Información requerida:</h4>
            <ul>
                <li>🆔 Datos del asegurado</li>
                <li>🛡️ Información del seguro</li>
                <li>💰 Costos y formas de pago</li>
                <li>📅 Fechas y vigencias</li>
                <li>📋 Coberturas y exclusiones</li>
                <li>📞 Información de contacto</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        st.info("""
        **💡 Tip:** Ten a mano la información del seguro para completar el formulario más rápidamente.
        """)

def show_interview():
    """Interfaz de entrevista conversacional"""
    
    st.header("💬 Conversación con tu asistente")
    
    # Sidebar con progreso
    with st.sidebar:
        st.header("📊 Progreso del certificado")
        
        total_questions = len(st.session_state.generator.variables)
        completed = len(st.session_state.responses)
        progress = completed / total_questions if total_questions > 0 else 0
        
        st.progress(progress)
        st.write(f"**Completado:** {completed}/{total_questions}")
        
        if st.session_state.responses:
            st.subheader("✅ Información guardada:")
            for var, response in st.session_state.responses.items():
                var_name = var.replace('_', ' ').title()
                display_response = str(response)[:30] + "..." if len(str(response)) > 30 else str(response)
                st.write(f"**{var_name}:** {display_response}")
        
        if st.button("🔄 Reiniciar"):
            st.session_state.clear()
            st.rerun()
    
    # Mostrar historial de chat
    for message in st.session_state.chat_history:
        if message["role"] == "assistant":
            st.chat_message("assistant", avatar="🤖").write(message["content"])
        else:
            st.chat_message("user", avatar="👤").write(message["content"])
    
    # Lógica de preguntas
    variables = st.session_state.generator.variables
    current_q = st.session_state.current_question
    
    if current_q < len(variables):
        current_var = variables[current_q]
        
        # Mostrar pregunta actual si no está en el historial
        if not st.session_state.chat_history or "?" not in st.session_state.chat_history[-1]["content"]:
            question_data = st.session_state.generator.questions[current_var]
            question = question_data['question']
            
            st.session_state.chat_history.append({
                "role": "assistant",
                "content": question
            })
            st.chat_message("assistant", avatar="🤖").write(question)
        
        # Campo de entrada
        show_input_field(current_var, current_q)
    
    else:
        # Todas las preguntas completadas
        st.session_state.chat_history.append({
            "role": "assistant", 
            "content": "🎉 ¡Excelente! He recopilado toda la información necesaria para generar tu certificado de Pacífico Seguros. El documento estará listo en un momento."
        })
        st.session_state.step = "generate"
        st.rerun()

def show_input_field(current_var, current_q):
    """Muestra el campo de entrada apropiado"""
    
    question_data = st.session_state.generator.questions[current_var]
    input_type = question_data['type']
    help_text = question_data['help']
    
    col1, col2 = st.columns([4, 1])
    
    with col1:
        if input_type == 'date':
            user_input = st.date_input(
                "Selecciona la fecha:",
                key=f"input_{current_q}",
                help=help_text
            )
            user_input = user_input.strftime("%d/%m/%Y")
        
        elif input_type == 'number':
            user_input = st.number_input(
                "Ingresa el monto:",
                min_value=0.0,
                step=0.01,
                key=f"input_{current_q}",
                help=help_text
            )
        
        elif input_type == 'email':
            user_input = st.text_input(
                "Escribe el email:",
                placeholder="ejemplo@correo.com",
                key=f"input_{current_q}",
                help=help_text
            )
        
        elif input_type == 'textarea':
            user_input = st.text_area(
                "Descripción detallada:",
                height=100,
                key=f"input_{current_q}",
                help=help_text
            )
        
        else:
            user_input = st.text_input(
                "Tu respuesta:",
                key=f"input_{current_q}",
                help=help_text
            )
    
    with col2:
        if st.button("Enviar", key=f"send_{current_q}"):
            handle_response(user_input, current_var, input_type)

def handle_response(user_input, current_var, input_type):
    """Procesa la respuesta del usuario"""
    
    # Validar respuesta
    if not user_input or str(user_input).strip() == "":
        st.error("Por favor proporciona una respuesta válida")
        return
    
    if input_type == 'email' and '@' not in str(user_input):
        st.error("Por favor ingresa un email válido")
        return
    
    # Agregar respuesta del usuario al chat
    st.session_state.chat_history.append({
        "role": "user",
        "content": str(user_input)
    })
    
    # Guardar respuesta
    st.session_state.responses[current_var] = user_input
    
    # Generar confirmación
    confirmations = [
        f"✅ Perfecto, guardé '{user_input}'",
        f"📝 Excelente, tengo '{user_input}' registrado",
        f"👍 Muy bien, información guardada",
        f"✨ Entendido, siguiente pregunta..."
    ]
    
    confirmation = random.choice(confirmations)
    st.session_state.chat_history.append({
        "role": "assistant",
        "content": confirmation
    })
    
    # Avanzar pregunta
    st.session_state.current_question += 1
    
    st.rerun()

def show_generation():
    """Pantalla de generación del certificado"""
    
    st.header("🎉 ¡Certificado listo para generar!")
    
    # Mostrar últimos mensajes del chat
    for message in st.session_state.chat_history[-2:]:
        if message["role"] == "assistant":
            st.chat_message("assistant", avatar="🤖").write(message["content"])
    
    # Resumen de información
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("📋 Información del certificado:")
        
        # Mostrar datos principales organizados
        st.markdown("**👤 Datos del Asegurado:**")
        st.write(f"• Nombre: {st.session_state.responses.get('nombre_asegurado', 'N/A')}")
        st.write(f"• Documento: {st.session_state.responses.get('tipo_documento', 'N/A')} - {st.session_state.responses.get('numero_documento', 'N/A')}")
        st.write(f"• Email: {st.session_state.responses.get('correo_asegurado', 'N/A')}")
        
        st.markdown("**🛡️ Información del Seguro:**")
        st.write(f"• Tipo: {st.session_state.responses.get('nombre_seguro', 'N/A')}")
        st.write(f"• Certificado: {st.session_state.responses.get('numero_certificado', 'N/A')}")
        st.write(f"• Póliza: {st.session_state.responses.get('numero_poliza', 'N/A')}")
        
        st.markdown("**💰 Información Financiera:**")
        st.write(f"• Costo: {st.session_state.responses.get('costo_total', 'N/A')} {st.session_state.responses.get('moneda', '')}")
        st.write(f"• Frecuencia: {st.session_state.responses.get('frecuencia_pago', 'N/A')}")
    
    with col2:
        st.metric(
            label="Campos completados",
            value=len(st.session_state.responses),
            delta=f"{len(st.session_state.generator.variables)} total"
        )
        
        st.info(f"""
        📄 **Tipo:** Certificado Pacífico Seguros
        
        ⏱️ **Generado:** {datetime.now().strftime('%H:%M')}
        
        ✅ **Estado:** Listo para descargar
        """)
    
    # Generar certificado
    if st.button("📄 Generar Certificado Pacífico Seguros", use_container_width=True):
        generate_certificate()
    
    # Opciones adicionales
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("🔄 Generar otro certificado", use_container_width=True):
            st.session_state.clear()
            st.rerun()
    
    with col2:
        if st.button("✏️ Modificar información", use_container_width=True):
            st.session_state.step = "interview"
            st.session_state.current_question = 0
            st.session_state.responses = {}
            st.session_state.chat_history = st.session_state.chat_history[:1]
            st.rerun()

def generate_certificate():
    """Genera y descarga el certificado"""
    
    with st.spinner("🤖 Generando tu certificado Pacífico Seguros..."):
        try:
            generator = st.session_state.generator
            
            # Generar documento
            doc = generator.generate_certificate_document(st.session_state.responses)
            
            # Guardar en memoria
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Nombre del archivo
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            nombre_asegurado = st.session_state.responses.get('nombre_asegurado', 'Asegurado')
            nombre_archivo = f"Certificado_Pacifico_{nombre_asegurado.replace(' ', '_')}_{timestamp}.docx"
            
            st.success("✅ ¡Certificado generado exitosamente!")
            
            # Botón de descarga
            st.download_button(
                label="⬇️ Descargar Certificado Pacífico Seguros",
                data=doc_buffer.getvalue(),
                file_name=nombre_archivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            # Mensaje final
            st.chat_message("assistant", avatar="🤖").write(
                "🎉 ¡Perfecto! Tu certificado de Pacífico Seguros está listo. "
                "El documento mantiene el formato oficial y tiene toda la información completada. "
                "¡Gracias por usar el generador!"
            )
            
            # Vista previa de información
            with st.expander("📋 Vista previa del certificado generado"):
                st.markdown(f"""
                **Certificado N° {st.session_state.responses.get('numero_certificado')}**
                
                **Asegurado:** {st.session_state.responses.get('nombre_asegurado')}
                
                **Seguro:** {st.session_state.responses.get('nombre_seguro')}
                
                **Póliza:** {st.session_state.responses.get('numero_poliza')}
                
                **Vigencia:** Desde {st.session_state.responses.get('fecha_inicio')}
                
                **Contacto:** {st.session_state.responses.get('correo_asegurado')}
                """)
            
        except Exception as e:
            st.error(f"❌ Error al generar el certificado: {str(e)}")
            st.info("Por favor verifica que toda la información esté completa e inténtalo de nuevo.")

# ====== EJECUTAR APLICACIÓN ======
if __name__ == "__main__":
    main()
